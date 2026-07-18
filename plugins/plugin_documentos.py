"""
plugin_documentos.py
=====================
Geracao e manipulacao de documentos: DOCX, XLSX, CSV,
arquivos de configuracao e templates.
"""

import os
import csv
import json
import io
import logging
from datetime import datetime

__version__ = "1.0.0"
PLUGIN_NAME = "Geracao de Documentos"
DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "agente_data", "documentos")
os.makedirs(DATA_DIR, exist_ok=True)


def register(api):
    def criar_csv(caminho: str, cabecalho: str, linhas: str, delimiter: str = ",") -> str:
        """Cria arquivo CSV. cabecalho: 'nome,idade' e linhas: '[["Joao",30],["Maria",25]]' (JSON)."""
        try:
            headers = [h.strip() for h in cabecalho.split(delimiter)]
            data = json.loads(linhas) if isinstance(linhas, str) else linhas
            parent = os.path.dirname(os.path.abspath(caminho))
            if parent:
                os.makedirs(parent, exist_ok=True)
            with open(caminho, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f, delimiter=delimiter)
                writer.writerow(headers)
                writer.writerows(data)
            return f"CSV criado: {caminho} ({len(data)} linhas, {len(headers)} colunas)."
        except Exception as e:
            return f"Erro: {e}"

    def criar_docx(caminho: str, titulo: str = "", conteudo: str = "") -> str:
        """Cria arquivo DOCX simples com titulo e paragrafos."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            return "Instale: pip install python-docx"
        try:
            parent = os.path.dirname(os.path.abspath(caminho))
            if parent:
                os.makedirs(parent, exist_ok=True)
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            if titulo:
                doc.add_heading(titulo, 0)
            for par in conteudo.split("\n"):
                par = par.strip()
                if par.startswith("# "):
                    doc.add_heading(par[2:], 1)
                elif par.startswith("## "):
                    doc.add_heading(par[3:], 2)
                elif par.startswith("### "):
                    doc.add_heading(par[4:], 3)
                elif par.startswith("- ") or par.startswith("* "):
                    doc.add_paragraph(par[2:], style="List Bullet")
                elif par.startswith("1. ") or par.startswith("2. "):
                    doc.add_paragraph(par[3:], style="List Number")
                elif par.strip():
                    doc.add_paragraph(par)
            doc.save(caminho)
            size = os.path.getsize(caminho)
            return f"DOCX criado: {caminho} ({size:,} bytes)."
        except Exception as e:
            return f"Erro: {e}"

    def criar_xlsx(caminho: str, dados: str, sheet_name: str = "Planilha1") -> str:
        """Cria arquivo XLSX. dados: JSON de objetos '[{"col": "val"}, ...]'."""
        try:
            import openpyxl
        except ImportError:
            return "Instale: pip install openpyxl"
        try:
            data = json.loads(dados) if isinstance(dados, str) else dados
            if not data:
                return "Dados vazios."
            parent = os.path.dirname(os.path.abspath(caminho))
            if parent:
                os.makedirs(parent, exist_ok=True)
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = sheet_name[:31]
            headers = list(data[0].keys())
            ws.append(headers)
            for item in data:
                ws.append([item.get(h, "") for h in headers])
            for col in ws.columns:
                max_len = max((len(str(cell.value or "")) for cell in col), default=0)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)
            wb.save(caminho)
            return f"XLSX criado: {caminho} ({len(data)} linhas, {len(headers)} colunas)."
        except Exception as e:
            return f"Erro: {e}"

    def xlsx_para_json(xlsx_path: str, sheet: str = "") -> str:
        """Converte planilha XLSX para JSON."""
        try:
            import openpyxl
        except ImportError:
            return "Instale: pip install openpyxl"
        try:
            wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
            ws = wb[sheet] if sheet else wb.active
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                return "Planilha vazia."
            headers = [str(h) if h else f"Coluna{i}" for i, h in enumerate(rows[0])]
            data = []
            for row in rows[1:]:
                item = {}
                for i, val in enumerate(row):
                    if i < len(headers):
                        item[headers[i]] = val
                data.append(item)
            wb.close()
            return json.dumps(data, ensure_ascii=False, indent=2, default=str)
        except Exception as e:
            return f"Erro: {e}"

    def docx_para_texto(docx_path: str) -> str:
        """Extrai texto de arquivo DOCX."""
        try:
            from docx import Document
        except ImportError:
            return "Instale: pip install python-docx"
        try:
            doc = Document(docx_path)
            paragraphs = []
            for p in doc.paragraphs:
                paragraphs.append(p.text)
            return "\n".join(paragraphs) or "Nenhum texto encontrado."
        except Exception as e:
            return f"Erro: {e}"

    def criar_template(nome: str, conteudo: str, pasta: str = "") -> str:
        """Salva um template de documento para reuso. Use {{variavel}} para placeholders."""
        try:
            template_dir = pasta or os.path.join(DATA_DIR, "templates")
            os.makedirs(template_dir, exist_ok=True)
            path = os.path.join(template_dir, nome)
            if not path.endswith((".txt", ".md", ".html", ".json", ".csv", ".py", ".js")):
                path += ".txt"
            with open(path, "w", encoding="utf-8") as f:
                f.write(conteudo)
            return f"Template salvo: {path}"
        except Exception as e:
            return f"Erro: {e}"

    def aplicar_template(template_path: str, variaveis: str) -> str:
        """Aplica variaveis a um template. variaveis: JSON '{"nome": "Joao", "idade": 30}'."""
        try:
            with open(template_path, "r", encoding="utf-8") as f:
                template = f.read()
            vars_dict = json.loads(variaveis) if isinstance(variaveis, str) else variaveis
            resultado = template
            for k, v in vars_dict.items():
                resultado = resultado.replace("{{" + k + "}}", str(v))
            return resultado
        except Exception as e:
            return f"Erro: {e}"

    def csv_resumir(csv_path: str, delimiter: str = ",") -> str:
        """Resumo estatistico de arquivo CSV: linhas, colunas, valores unicos por coluna."""
        try:
            with open(csv_path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                rows = list(reader)
            if not rows:
                return "CSV vazio."
            headers = reader.fieldnames
            lines = [f"Arquivo: {csv_path}", f"Linhas: {len(rows)}", f"Colunas: {len(headers)}", ""]
            for h in headers:
                vals = [r[h] for r in rows if r.get(h)]
                unique = len(set(vals))
                non_empty = sum(1 for v in vals if v.strip())
                lines.append(f"  {h}: {unique} valores unicos, {non_empty}/{len(rows)} preenchidos")
            return "\n".join(lines)
        except Exception as e:
            return f"Erro: {e}"

    api.register_tool("criar_csv", criar_csv,
        "Cria arquivo CSV. cabecalho: 'nome,idade'. linhas: JSON array de arrays.",
        {"caminho": {"type": "string", "description": "Caminho do arquivo"}, "cabecalho": {"type": "string", "description": "Cabecalho separado por virgula"}, "linhas": {"type": "string", "description": "JSON array de arrays de valores"}, "delimiter": {"type": "string", "description": "Delimitador (opcional)"}}, ["caminho", "cabecalho", "linhas"])

    api.register_tool("criar_docx", criar_docx,
        "Cria arquivo DOCX com titulo e paragrafos formatados (suporta # ## ### - *).",
        {"caminho": {"type": "string", "description": "Caminho do .docx"}, "titulo": {"type": "string", "description": "Titulo do documento (opcional)"}, "conteudo": {"type": "string", "description": "Conteudo em texto simples com marcadores"}}, ["caminho", "conteudo"])

    api.register_tool("criar_xlsx", criar_xlsx,
        "Cria arquivo XLSX a partir de JSON array de objetos.",
        {"caminho": {"type": "string", "description": "Caminho do .xlsx"}, "dados": {"type": "string", "description": "JSON array de objetos"}, "sheet_name": {"type": "string", "description": "Nome da planilha (opcional)"}}, ["caminho", "dados"])

    api.register_tool("xlsx_para_json", xlsx_para_json,
        "Converte planilha XLSX para JSON.",
        {"xlsx_path": {"type": "string", "description": "Caminho do .xlsx"}, "sheet": {"type": "string", "description": "Nome da planilha (opcional)"}}, ["xlsx_path"])

    api.register_tool("docx_para_texto", docx_para_texto,
        "Extrai texto de arquivo DOCX.",
        {"docx_path": {"type": "string", "description": "Caminho do .docx"}}, ["docx_path"])

    api.register_tool("criar_template", criar_template,
        "Salva template com placeholders {{variavel}} para reuso.",
        {"nome": {"type": "string", "description": "Nome do template"}, "conteudo": {"type": "string", "description": "Conteudo com {{variaveis}}"}, "pasta": {"type": "string", "description": "Pasta para salvar (opcional)"}}, ["nome", "conteudo"])

    api.register_tool("aplicar_template", aplicar_template,
        "Aplica variaveis JSON a um template {{variavel}}.",
        {"template_path": {"type": "string", "description": "Caminho do template"}, "variaveis": {"type": "string", "description": "JSON com variaveis"}}, ["template_path", "variaveis"])

    api.register_tool("csv_resumir", csv_resumir,
        "Resumo estatistico de CSV: colunas, valores unicos, preenchimento.",
        {"csv_path": {"type": "string", "description": "Caminho do CSV"}, "delimiter": {"type": "string", "description": "Delimitador (opcional)"}}, ["csv_path"])

    return {
        "name": PLUGIN_NAME,
        "version": __version__,
        "description": "Geracao de documentos: CSV, DOCX, XLSX, templates e conversoes",
        "tools": ["criar_csv", "criar_docx", "criar_xlsx", "xlsx_para_json", "docx_para_texto", "criar_template", "aplicar_template", "csv_resumir"],
    }
